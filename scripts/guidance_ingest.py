#!/usr/bin/env python3
"""
Ingest local best-practice / guidance documents into the brain.

WHY THIS EXISTS
    The brain had no ABAP Cloud or RAP guidance at all. Measured 2026-09-04, a core
    query ("RAP business object behaviour definition managed implementation") scored
    0.34-0.41 across every hit -- uniformly weak, because nothing in the corpus
    covered it. The UI5/CAP/Node harvest fixed side-by-side grounding and left
    developer extensibility, the mode most S/4HANA Public Cloud work actually uses,
    unsupported.

    Vendor docs come from webdocs_ingest.py. This handles the other kind: internal
    review checklists and delivery standards that exist as a file on someone's disk,
    not at a URL.

WHY IT IS NOT webdocs_ingest.py
    That script harvests public vendor documentation over HTTP. These documents are
    internal, often carry a client name, and are PII-masked on the way in. Keeping
    them apart keeps the provenance distinction visible instead of implicit.

PROVENANCE -- read before adding a source
    Tagged source_system="abap_guidance", which is deliberately:
      * IN  brain_regression.PROVENANCE_EXEMPT_SOURCES  -- guidance is
        phase-independent, so a phase filter must not hide it, and
      * OUT of brain_regression.PUBLIC_SOURCE_SYSTEMS   -- these are NOT public by
        construction, so document names stay hashed in the committed baseline.
    Text is masked with sharepoint_ingest.mask() and brain/ is gitignored, so the
    content never reaches git either way.

USAGE
    python3.11 scripts/guidance_ingest.py                    # ingest brain/guidance/raw/
    python3.11 scripts/guidance_ingest.py --dry-run          # report, write nothing
    Then: python3.11 scripts/embed_chunks.py && pm2 restart s4pc-mcp

    Drop .docx / .pdf / .md / .txt files into brain/guidance/raw/ first.
"""

import argparse
import hashlib
import json
import sys
import time
from pathlib import Path

sys.path.insert(0, str(Path(__file__).resolve().parent))

BASE_DIR   = Path(__file__).resolve().parent.parent
OUT_DIR    = BASE_DIR / "brain" / "guidance"
RAW_DIR    = OUT_DIR / "raw"
CHUNKS_DIR = OUT_DIR / "chunks"
MANIFEST   = OUT_DIR / "manifest.json"

SUPPORTED = {".docx", ".pdf", ".md", ".txt"}
MIN_USEFUL_CHARS = 400        # a guidance doc shorter than this is a stub or a bad read


def extract(path):
    """Plain text out of a guidance document. Returns '' if unreadable."""
    ext = path.suffix.lower()
    if ext in (".md", ".txt"):
        return path.read_text(encoding="utf-8", errors="replace")
    if ext == ".docx":
        try:
            import docx
        except ImportError:
            sys.exit("python-docx missing. Run: pip3.11 install python-docx")
        d = docx.Document(str(path))
        parts = [p.text for p in d.paragraphs if p.text.strip()]
        # Checklists and rule tables are the substance of a review document, so the
        # tables matter at least as much as the paragraphs -- python-docx keeps them
        # out of .paragraphs entirely and they would otherwise be dropped silently.
        for t in d.tables:
            for row in t.rows:
                cells = [c.text.strip() for c in row.cells if c.text.strip()]
                if cells:
                    parts.append(" | ".join(cells))
        return "\n".join(parts)
    if ext == ".pdf":
        try:
            import fitz
        except ImportError:
            sys.exit("pymupdf missing. Run: pip3.11 install pymupdf")
        with fitz.open(str(path)) as doc:
            return "\n".join(page.get_text() for page in doc)
    return ""


def main():
    ap = argparse.ArgumentParser(description="Ingest local guidance documents")
    ap.add_argument("--dry-run", action="store_true", help="report only, write nothing")
    ap.add_argument("--no-mask", action="store_true",
                    help="Skip PII masking. Only for content you have confirmed is public.")
    args = ap.parse_args()

    if not RAW_DIR.exists() or not any(RAW_DIR.iterdir()):
        sys.exit("No files in %s — copy guidance documents there first." % RAW_DIR)

    from sharepoint_ingest import chunk, mask     # reuse the corpus's own chunker/masker

    files = [f for f in sorted(RAW_DIR.rglob("*"))
             if f.is_file() and f.suffix.lower() in SUPPORTED]
    print("== %d guidance document(s)%s" % (len(files), " (DRY RUN)" if args.dry_run else ""))

    manifest, n_ok, n_bad, n_chunks = [], 0, 0, 0
    for f in files:
        rel = f.relative_to(RAW_DIR).as_posix()
        try:
            text = extract(f)
        except Exception as exc:                  # a bad file must not kill the batch
            print("   FAILED  %s  (%s)" % (rel, str(exc)[:70]))
            manifest.append({"file": rel, "status": "FAILED", "error": str(exc)[:160]})
            n_bad += 1
            continue

        if len(text) < MIN_USEFUL_CHARS:
            print("   THIN    %s  (%d chars — not stored)" % (rel, len(text)))
            manifest.append({"file": rel, "status": "THIN", "chars": len(text)})
            n_bad += 1
            continue

        if not args.no_mask:
            text = mask(text)
        pieces = chunk(text)
        did = hashlib.sha1(rel.encode("utf-8")).hexdigest()[:12]
        print("   ok      %s  %d chars -> %d chunks" % (rel, len(text), len(pieces)))
        n_ok += 1
        n_chunks += len(pieces)

        if not args.dry_run:
            out = CHUNKS_DIR
            out.mkdir(parents=True, exist_ok=True)
            for i, piece in enumerate(pieces):
                (out / ("%s_%04d.json" % (did, i))).write_text(json.dumps({
                    "id":               "%s_%04d" % (did, i),
                    "text":             piece,
                    "source":           f.stem,
                    "source_system":    "abap_guidance",
                    "deliverable_type": "abap_guidance",
                    "content_type":     "reference",
                    # phase/agent_role are delivery provenance and guidance has none.
                    # They are set only because every consumer's filters expect the
                    # fields to exist; PROVENANCE_EXEMPT_SOURCES is what stops a phase
                    # filter from hiding these. See scripts/vectorstore.py.
                    "phase":            "Reference",
                    "agent_role":       "reference",
                    "relative_path":    rel,
                }, ensure_ascii=False), encoding="utf-8")
        manifest.append({"file": rel, "status": "ok",
                         "chars": len(text), "chunks": len(pieces),
                         "masked": not args.no_mask})

    if not args.dry_run:
        OUT_DIR.mkdir(parents=True, exist_ok=True)
        MANIFEST.write_text(json.dumps({
            "ingested_at": time.strftime("%Y-%m-%dT%H:%M:%SZ", time.gmtime()),
            "documents": manifest,
        }, indent=2, ensure_ascii=False), encoding="utf-8")

    print("\n== %d ok (%d chunks), %d skipped" % (n_ok, n_chunks, n_bad))
    if n_ok and not args.dry_run:
        print("== Next: python3.11 scripts/embed_chunks.py   (then pm2 restart s4pc-mcp)")
    sys.exit(0 if n_ok else 1)


if __name__ == "__main__":
    main()
